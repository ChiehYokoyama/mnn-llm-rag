"""
文档管理器模块 - 支持多种文件格式的加载和处理
支持格式：TXT, Markdown, DOCX, PDF
"""
import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from abc import ABC, abstractmethod
import json
from datetime import datetime

logger = logging.getLogger(__name__)


# ============================================================================
# 文档加载器基类
# ============================================================================

class DocumentLoader(ABC):
    """文档加载器基类"""

    @abstractmethod
    def load(self, file_path: str) -> List[str]:
        """
        加载文档并返回文本内容

        Args:
            file_path: 文件路径

        Returns:
            文本列表，每个元素是一个段落
        """
        pass

    @staticmethod
    def get_file_info(file_path: str) -> Dict:
        """获取文件信息"""
        path = Path(file_path)
        return {
            'name': path.name,
            'size': path.stat().st_size,
            'created': datetime.fromtimestamp(path.stat().st_ctime).isoformat(),
            'modified': datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
        }


# ============================================================================
# TXT 文档加载器
# ============================================================================

class TxtLoader(DocumentLoader):
    """文本文档加载器"""

    def load(self, file_path: str) -> List[str]:
        """加载TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # 按行分割，过滤空行
            paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            logger.info(f"✅ TXT文件加载成功: {len(paragraphs)} 段")
            return paragraphs

        except Exception as e:
            logger.error(f"❌ TXT文件加载失败: {e}")
            return []


# ============================================================================
# Markdown 文档加载器
# ============================================================================

class MarkdownLoader(DocumentLoader):
    """Markdown文档加载器"""

    def load(self, file_path: str) -> List[str]:
        """加载Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            paragraphs = self._parse_markdown(content)
            logger.info(f"✅ Markdown文件加载成功: {len(paragraphs)} 段")
            return paragraphs

        except Exception as e:
            logger.error(f"❌ Markdown文件加载失败: {e}")
            return []

    @staticmethod
    def _parse_markdown(content: str) -> List[str]:
        """解析Markdown内容"""
        # 移除Markdown特殊字符但保留内容
        content = re.sub(r'#+\s+', '', content)  # 移除标题符号
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)  # 转换链接
        content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)  # 移除粗体
        content = re.sub(r'__([^_]+)__', r'\1', content)  # 移除粗体
        content = re.sub(r'\*([^*]+)\*', r'\1', content)  # 移除斜体
        content = re.sub(r'_([^_]+)_', r'\1', content)  # 移除斜体
        content = re.sub(r'`([^`]+)`', r'\1', content)  # 移除代码块标记
        content = re.sub(r'```[\s\S]*?```', '', content)  # 移除代码块

        # 按段落分割
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        return paragraphs


# ============================================================================
# DOCX 文档加载器
# ============================================================================

class DocxLoader(DocumentLoader):
    """DOCX文档加载器"""

    def load(self, file_path: str) -> List[str]:
        """加载DOCX文件"""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        paragraphs.append(row_text)

            logger.info(f"✅ DOCX文件加载成功: {len(paragraphs)} 段")
            return paragraphs

        except ImportError:
            logger.error("❌ 缺少python-docx库，请运行: pip install python-docx")
            return []
        except Exception as e:
            logger.error(f"❌ DOCX文件加载失败: {e}")
            return []


# ============================================================================
# PDF 文档加载器
# ============================================================================

class PdfLoader(DocumentLoader):
    """PDF文档加载器"""

    def load(self, file_path: str) -> List[str]:
        """加载PDF文件"""
        try:
            import pypdf

            paragraphs = []

            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()

                    if text.strip():
                        # 按行分割，过滤空行
                        lines = [line.strip() for line in text.split('\n') if line.strip()]
                        paragraphs.extend(lines)

            logger.info(f"✅ PDF文件加载成功: {len(paragraphs)} 段")
            return paragraphs

        except ImportError:
            logger.error("❌ 缺少pypdf库，请运行: pip install pypdf")
            return []
        except Exception as e:
            logger.error(f"❌ PDF文件加载失败: {e}")
            return []


# ============================================================================
# 文档管理器
# ============================================================================

class DocumentManager:
    """文档管理器 - 支持多种格式的文档加载和处理"""

    # 支持的文件格式
    SUPPORTED_FORMATS = {
        '.txt': TxtLoader,
        '.md': MarkdownLoader,
        '.markdown': MarkdownLoader,
        '.docx': DocxLoader,
        '.doc': DocxLoader,
        '.pdf': PdfLoader,
    }

    def __init__(self, chunk_size: int = 200, overlap: int = 50):
        """
        初始化文档管理器

        Args:
            chunk_size: 分块大小（字符数）
            overlap: 分块重叠字符数
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.documents = {}  # 存储已加载的文档
        self.file_metadata = {}  # 存储文件元数据

    def load_document(self, file_path: str) -> Tuple[bool, str, List[str]]:
        """
        加载单个文档

        Args:
            file_path: 文件路径

        Returns:
            (是否成功, 消息, 文本列表)
        """
        file_path = str(file_path)

        # 检查文件是否存在
        if not os.path.exists(file_path):
            msg = f"文件不存在: {file_path}"
            logger.error(f"❌ {msg}")
            return False, msg, []

        # 获取文件扩展名
        ext = Path(file_path).suffix.lower()

        # 检查是否支持该格式
        if ext not in self.SUPPORTED_FORMATS:
            supported = ', '.join(self.SUPPORTED_FORMATS.keys())
            msg = f"不支持的文件格式: {ext}。支持的格式: {supported}"
            logger.error(f"❌ {msg}")
            return False, msg, []

        # 加载文档
        loader = self.SUPPORTED_FORMATS[ext]()
        paragraphs = loader.load(file_path)

        if not paragraphs:
            msg = f"文件加载失败或内容为空: {file_path}"
            logger.error(f"❌ {msg}")
            return False, msg, []

        # 存储文档和元数据
        doc_name = Path(file_path).stem
        self.documents[doc_name] = paragraphs
        self.file_metadata[doc_name] = {
            'path': file_path,
            'format': ext,
            'file_info': loader.get_file_info(file_path),
            'paragraph_count': len(paragraphs),
            'total_chars': sum(len(p) for p in paragraphs),
            'loaded_at': datetime.now().isoformat(),
        }

        msg = f"文档加载成功: {len(paragraphs)} 段，总计 {self.file_metadata[doc_name]['total_chars']} 字符"
        logger.info(f"✅ {msg}")
        return True, msg, paragraphs

    def load_documents_from_directory(self, directory: str) -> Dict[str, Tuple[bool, str, List[str]]]:
        """
        从目录加载所有支持的文档

        Args:
            directory: 目录路径

        Returns:
            {文档名: (是否成功, 消息, 文本列表)}
        """
        results = {}

        if not os.path.isdir(directory):
            logger.error(f"❌ 目录不存在: {directory}")
            return results

        # 查找所有支持的文件
        for ext in self.SUPPORTED_FORMATS.keys():
            for file_path in Path(directory).glob(f'*{ext}'):
                success, msg, paragraphs = self.load_document(str(file_path))
                doc_name = file_path.stem
                results[doc_name] = (success, msg, paragraphs)

        logger.info(f"📁 目录加载完成: {len(results)} 个文件")
        return results

    def chunk_text(self, text: str, chunk_size: Optional[int] = None,
                   overlap: Optional[int] = None) -> List[str]:
        """
        将文本分块

        Args:
            text: 输入文本
            chunk_size: 分块大小，不指定则使用默认值
            overlap: 重叠大小，不指定则使用默认值

        Returns:
            分块列表
        """
        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.overlap

        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunks.append(text[start:end])
            start = end - overlap if end < len(text) else end

        return chunks

    def process_documents(self, chunk_size: Optional[int] = None) -> List[str]:
        """
        处理所有已加载的文档，进行分块和清理

        Args:
            chunk_size: 分块大小

        Returns:
            处理后的文本分块列表
        """
        all_chunks = []

        for doc_name, paragraphs in self.documents.items():
            logger.info(f"处理文档: {doc_name}")

            for para in paragraphs:
                # 清理文本
                cleaned = self._clean_text(para)
                if cleaned:
                    # 分块
                    chunks = self.chunk_text(cleaned, chunk_size)
                    all_chunks.extend(chunks)

        logger.info(f"✅ 文档处理完成: 生成 {len(all_chunks)} 个文本块")
        return all_chunks

    @staticmethod
    def _clean_text(self, text: str) -> str:
        """清理文本并去重"""
        # 移除多余空格和特殊符号，保留常用的汉字、英文字母、数字和标点符号
        text = re.sub(r'\s+', ' ', text)  # 替换多余空白为空格
        text = re.sub(
            r'[^\u4e00-\u9fa5a-zA-Z0-9\s，。！？；：、\'"\'""''（）【】《》]', '', text
        )  # 允许的字符范围
        return text.strip()

    def get_document_stats(self) -> Dict:
        """获取文档统计信息"""
        stats = {
            'total_documents': len(self.documents),
            'total_paragraphs': sum(len(p) for p in self.documents.values()),
            'total_characters': sum(
                self.file_metadata[doc_name]['total_chars']
                for doc_name in self.documents.keys()
            ),
            'documents': {}
        }

        for doc_name, metadata in self.file_metadata.items():
            stats['documents'][doc_name] = {
                'format': metadata['format'],
                'paragraphs': metadata['paragraph_count'],
                'characters': metadata['total_chars'],
                'path': metadata['path'],
            }

        return stats

    def export_metadata(self, output_path: str):
        """
        导出文件元数据为JSON

        Args:
            output_path: 输出文件路径
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(self.file_metadata, f, ensure_ascii=False, indent=2)
            logger.info(f"✅ 元数据已导出: {output_path}")
        except Exception as e:
            logger.error(f"❌ 元数据导出失败: {e}")

    def clear_documents(self):
        """清空所有已加载的文档"""
        self.documents.clear()
        self.file_metadata.clear()
        logger.info("✅ 所有文档已清空")

    def get_loaded_documents_count(self) -> int:
        """获取已加载的文档数量"""
        return len(self.documents)

    def _remove_duplicates(self, fragments: List[str]) -> List[str]:
        """移除重复片段"""
        seen = set()  # 存储已见片段
        unique_fragments = []
        for frag in fragments:
            if frag not in seen:  # 去重逻辑
                seen.add(frag)
                unique_fragments.append(frag)
        logger.info(f"⚙️ 去重完成，共移除 {len(fragments) - len(unique_fragments)} 个重复片段")
        return unique_fragments